#!/usr/bin/env python3
"""Wandelt ein (Teilmengen-)Markdown in eine echte .docx um - ohne Microsoft Word.

Aufruf:
    python build_docx.py <input.md> <output.docx> [--light] [--no-toc]

Optionen:
    --light    helles Code-Theme (druckfreundlich) statt des dunklen Bildschirm-Themes
    --no-toc   kein Inhaltsverzeichnis einfuegen

Unterstuetztes Markdown:
    # .. ####     Ueberschriften (Ebene 1-4)
    Absaetze      Fliesstext mit **fett** und `code`
    - / *         Aufzaehlung
    | a | b |     Tabellen (erste Zeile = Kopf, Trennzeile ---|--- wird ignoriert)
    ```lang       Codebloecke mit Syntax-Highlighting (Java, Python, SQL, JSON, XML, JS)
    >             Zitat / "Merke"-Kasten (Rahmen links)
    ---           Trennlinie

Exit-Codes: 0 = ok, 1 = falscher Aufruf, 3 = python-docx fehlt.
"""
import sys
import re
import pathlib

# UTF-8-Ausgabe erzwingen (Windows-Konsole ist oft cp1252 -> Mojibake)
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except AttributeError:
    pass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("DOCX_MISSING: 'python-docx' ist nicht installiert. "
                     "Bitte 'pip install --user python-docx' ausfuehren.\n")
    sys.exit(3)


# ---------------------------------------------------------------------------
# Regex-Bausteine
# ---------------------------------------------------------------------------
INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')
SEP_ROW_RE = re.compile(r'^\s*\|[\s:\-\|]+\|?\s*$')
HEADING_RE = re.compile(r'^(#{1,4})\s+(.*)')
QUOTE_RE = re.compile(r'^>\s?')
BULLET_RE = re.compile(r'^\s*[-*]\s+(.*)')
RULE_RE = re.compile(r'^\s*---\s*$')
WORD_RE = re.compile(r'@?\w+')
NUM_RE = re.compile(r'\d+(\.\d+)?')

# ---------------------------------------------------------------------------
# Farbthemen fuer Codebloecke
#   dark  = Bildschirm (fast schwarz, VS-Code-artig) - Standard
#   light = Druck (heller Hintergrund, spart Toner, gut auf S/W-Druckern)
# ---------------------------------------------------------------------------
THEMES = {
    "dark": {"BG": "1E1E1E", "BORDER": "3C3C3C", "TEXT": "D4D4D4",
             "KEYWORD": "569CD6", "STRING": "CE9178", "COMMENT": "6A9955",
             "NUMBER": "B5CEA8", "TAG": "569CD6"},
    "light": {"BG": "F6F8FA", "BORDER": "D0D7DE", "TEXT": "24292E",
              "KEYWORD": "0033B3", "STRING": "067D17", "COMMENT": "6A737D",
              "NUMBER": "1750EB", "TAG": "0033B3"},
}
C_LABEL = "6E7781"    # Sprach-Label (auf weisser Seite, themenunabhaengig)

# Aktive Code-Farben - werden von apply_theme() gesetzt.
C_BG = C_BORDER = C_TEXT = C_KEYWORD = C_STRING = C_COMMENT = C_NUMBER = C_TAG = ""


def apply_theme(name):
    global C_BG, C_BORDER, C_TEXT, C_KEYWORD, C_STRING, C_COMMENT, C_NUMBER, C_TAG
    t = THEMES.get(name, THEMES["dark"])
    C_BG, C_BORDER, C_TEXT = t["BG"], t["BORDER"], t["TEXT"]
    C_KEYWORD, C_STRING, C_COMMENT = t["KEYWORD"], t["STRING"], t["COMMENT"]
    C_NUMBER, C_TAG = t["NUMBER"], t["TAG"]


apply_theme("dark")

# Farben fuer Tabellen und Merke-Kaesten (themenunabhaengig)
T_HEADER = "2F5496"   # Tabellenkopf (dunkelblau)
T_BAND = "EAF0F8"     # Zebra-Streifen (hellblau)
T_BORDER = "C7D0E0"   # Tabellenrahmen (dezent)
Q_BG = "EEF3FB"       # Merke-Kasten Hintergrund
Q_BAR = "4472C4"      # Merke-Kasten Balken links

# ---------------------------------------------------------------------------
# Schluesselwoerter je Sprache
# ---------------------------------------------------------------------------
KW = {
    "java": set("""abstract assert boolean break byte case catch char class const
        continue default do double else enum extends final finally float for goto if
        implements import instanceof int interface long native new package private
        protected public return short static strictfp super switch synchronized this
        throw throws transient try void volatile while true false null var""".split()),
    "python": set("""def class return if elif else for while import from as try except
        finally with lambda None True False and or not in is pass break continue global
        nonlocal yield raise del assert async await elif print self""".split()),
    "sql": set("""select from where insert into values update set delete create table
        alter drop join inner left right outer full cross on group by order having
        primary key foreign references not null unique and or as count sum avg min max
        distinct limit offset like between in is asc desc auto_increment varchar int
        default union all exists""".split()),
    "json": set("true false null".split()),
    "js": set("""var let const function return if else for while do switch case break
        continue new class extends this typeof instanceof null true false undefined
        import export from default async await try catch finally throw of in""".split()),
}
KW["javascript"] = KW["js"]
KW["python3"] = KW["python"]

# Kommentar-Marker je Sprache: (Zeilenkommentar, Blockstart, Blockende)
COMMENTS = {
    "java": ("//", "/*", "*/"),
    "js": ("//", "/*", "*/"),
    "javascript": ("//", "/*", "*/"),
    "sql": ("--", "/*", "*/"),
    "python": ("#", None, None),
    "python3": ("#", None, None),
    "json": (None, None, None),
    "xml": (None, "<!--", "-->"),
    "html": (None, "<!--", "-->"),
}


def _kw_hit(word, lang):
    kws = KW.get(lang)
    if not kws:
        return False
    if lang == "sql":
        return word.lower() in kws
    return word in kws


def tokenize_code(line, lang, state):
    """Zerlegt eine Codezeile in (text, farbe|None)-Paare fuer das Highlighting."""
    out = []
    i, n = 0, len(line)
    line_c, block_o, block_c = COMMENTS.get(lang, (None, None, None))

    # XML/HTML: eigener, einfacher Modus (Tags + Strings + Kommentare)
    if lang in ("xml", "html"):
        return _tokenize_xml(line, state)

    while i < n:
        # laufender Blockkommentar
        if state.get("block"):
            end = line.find(block_c, i) if block_c else -1
            if end == -1:
                out.append((line[i:], C_COMMENT)); i = n
            else:
                out.append((line[i:end + len(block_c)], C_COMMENT))
                i = end + len(block_c); state["block"] = False
            continue
        # Blockkommentar-Beginn
        if block_o and line.startswith(block_o, i):
            end = line.find(block_c, i + len(block_o))
            if end == -1:
                out.append((line[i:], C_COMMENT)); i = n; state["block"] = True
            else:
                out.append((line[i:end + len(block_c)], C_COMMENT)); i = end + len(block_c)
            continue
        # Zeilenkommentar
        if line_c and line.startswith(line_c, i):
            out.append((line[i:], C_COMMENT)); i = n
            continue
        ch = line[i]
        # Zeichenkette
        if ch == '"' or ch == "'":
            j = i + 1
            while j < n and line[j] != ch:
                if line[j] == "\\":
                    j += 1
                j += 1
            j = min(j + 1, n)
            out.append((line[i:j], C_STRING)); i = j
            continue
        # Wort (Schluesselwort / Zahl / normal)
        m = WORD_RE.match(line, i)
        if m:
            w = m.group(0)
            if NUM_RE.fullmatch(w):
                color = C_NUMBER
            elif _kw_hit(w, lang):
                color = C_KEYWORD
            else:
                color = None
            out.append((w, color)); i = m.end()
            continue
        # sonstiges Einzelzeichen
        out.append((ch, None)); i += 1
    return out


def _tokenize_xml(line, state):
    out = []
    i, n = 0, len(line)
    while i < n:
        if state.get("block"):
            end = line.find("-->", i)
            if end == -1:
                out.append((line[i:], C_COMMENT)); i = n
            else:
                out.append((line[i:end + 3], C_COMMENT)); i = end + 3; state["block"] = False
            continue
        if line.startswith("<!--", i):
            end = line.find("-->", i + 4)
            if end == -1:
                out.append((line[i:], C_COMMENT)); i = n; state["block"] = True
            else:
                out.append((line[i:end + 3], C_COMMENT)); i = end + 3
            continue
        ch = line[i]
        if ch == '"' or ch == "'":
            j = i + 1
            while j < n and line[j] != ch:
                j += 1
            j = min(j + 1, n)
            out.append((line[i:j], C_STRING)); i = j
            continue
        if ch == "<":
            m = re.match(r'</?[\w:.-]+', line[i:])
            if m:
                out.append((m.group(0), C_TAG)); i += len(m.group(0)); continue
        if ch == ">" or (ch == "/" and i + 1 < n and line[i + 1] == ">"):
            out.append((ch, C_TAG)); i += 1
            continue
        m = WORD_RE.match(line, i)
        if m:
            out.append((m.group(0), None)); i = m.end()
            continue
        out.append((ch, None)); i += 1
    return out


# ---------------------------------------------------------------------------
# Inline-Formatierung fuer normalen Text
# ---------------------------------------------------------------------------
def add_inline(paragraph, text):
    """Fuegt Text mit **fett** und `code` als Runs hinzu."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# OOXML-Hilfsfunktionen (Schattierung, Rahmen)
# ---------------------------------------------------------------------------
def _shade(pPr, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def shade_paragraph(paragraph, fill):
    _shade(paragraph._p.get_or_add_pPr(), fill)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def add_border(paragraph, side, color="AAAAAA", sz="6", space="1"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    edge = OxmlElement('w:' + side)
    edge.set(qn('w:val'), 'single')
    edge.set(qn('w:sz'), sz)
    edge.set(qn('w:space'), space)
    edge.set(qn('w:color'), color)
    pBdr.append(edge)


def set_cell_border(cell, color, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def _preserve_space(run):
    t = run._r.find(qn('w:t'))
    if t is not None:
        t.set(qn('xml:space'), 'preserve')


def set_table_borders(table, color, sz="4"):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def set_table_cell_margins(table, top=40, bottom=40, left=110, right=110):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        e = OxmlElement('w:' + side)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tblPr.append(mar)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), 'true')
    trPr.append(th)


# ---------------------------------------------------------------------------
# Bloecke: Tabelle und Code-Box
# ---------------------------------------------------------------------------
def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    set_table_borders(table, T_BORDER, "4")
    set_table_cell_margins(table, top=50, bottom=50, left=120, right=120)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            val = row[ci] if ci < len(row) else ""
            if ri == 0:
                run = para.add_run(val)
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, T_HEADER)
            else:
                add_inline(para, val)
                for r in para.runs:
                    r.font.size = Pt(10)
                if ri % 2 == 0:  # jede zweite Datenzeile leicht einfaerben
                    shade_cell(cell, T_BAND)
    set_repeat_header(table.rows[0])
    doc.add_paragraph()


def add_code_block(doc, code_lines, lang):
    """Rendert einen Codeblock als Code-Box mit Syntax-Highlighting."""
    # Fuehrende/abschliessende Leerzeilen entfernen
    while code_lines and code_lines[0].strip() == "":
        code_lines.pop(0)
    while code_lines and code_lines[-1].strip() == "":
        code_lines.pop()
    if not code_lines:
        return

    lang = (lang or "").lower()

    # Optionales Sprach-Label ueber der Box
    if lang and lang not in ("text", "plain", "output"):
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(0)
        run = cap.add_run(lang.upper())
        run.bold = True
        run.font.size = Pt(7)
        run.font.name = "Consolas"
        run.font.color.rgb = RGBColor(0x6E, 0x77, 0x81)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    set_table_cell_margins(table, top=80, bottom=80, left=140, right=140)
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_border(cell, C_BORDER)
    shade_cell(cell, C_BG)

    state = {"block": False}
    first = True
    for line in code_lines:
        line = line.replace("\t", "    ")
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        if line.strip() == "":
            continue
        for text, color in tokenize_code(line, lang, state):
            run = para.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(color if color else C_TEXT)
            _preserve_space(run)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Dokument-Stil (Schrift, Ueberschriften, Seitenraender)
# ---------------------------------------------------------------------------
def style_document(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12

    heads = {1: ('1F3864', 17), 2: ('2E5496', 14), 3: ('2E5496', 12), 4: ('44546A', 11)}
    for lvl, (col, size) in heads.items():
        try:
            st = doc.styles['Heading %d' % lvl]
            st.font.color.rgb = RGBColor.from_string(col)
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.name = 'Calibri'
        except KeyError:
            pass

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def _add_field(paragraph, instr, placeholder=""):
    """Fuegt ein Word-Feld (z.B. TOC, PAGE) in einen Absatz ein."""
    r1 = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fld_begin)

    r2 = paragraph.add_run()
    instr_el = OxmlElement('w:instrText')
    instr_el.set(qn('xml:space'), 'preserve')
    instr_el.text = instr
    r2._r.append(instr_el)

    r3 = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fld_sep)

    if placeholder:
        paragraph.add_run(placeholder)

    r4 = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r4._r.append(fld_end)


def add_table_of_contents(doc):
    """Fuegt eine Ueberschrift + TOC-Feld ein (Word fuellt es beim Oeffnen)."""
    head = doc.add_paragraph()
    run = head.add_run("Inhaltsverzeichnis")
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor.from_string('1F3864')
    head.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    _add_field(p, ' TOC \\o "1-3" \\h \\z \\u ',
               "(Inhaltsverzeichnis: Rechtsklick > Felder aktualisieren bzw. F9)")
    doc.add_page_break()


def add_page_numbers(doc):
    """Zentrierte 'Seite X von Y'-Fusszeile in allen Abschnitten."""
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Seite ")
        _add_field(p, ' PAGE ')
        p.add_run(" von ")
        _add_field(p, ' NUMPAGES ')
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string('808080')


def enable_update_fields(doc):
    """Sorgt dafuer, dass Word Felder (TOC) beim Oeffnen aktualisiert."""
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        el = OxmlElement('w:updateFields')
        el.set(qn('w:val'), 'true')
        settings.append(el)


# ---------------------------------------------------------------------------
# Hauptparser
# ---------------------------------------------------------------------------
def build(md_path, docx_path, theme="dark", toc=True):
    apply_theme(theme)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_page_numbers(doc)

    in_code = False
    code_buf = []
    code_lang = ""
    table_rows = []
    toc_done = False

    def flush_table():
        if table_rows:
            add_table(doc, table_rows)
            table_rows.clear()

    for raw in lines:
        line = raw

        quoted = False
        if QUOTE_RE.match(line):
            quoted = True
            line = QUOTE_RE.sub('', line, count=1)

        # Codeblock-Grenzen
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, code_buf, code_lang)
                code_buf = []
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Tabellenzeilen
        if line.strip().startswith("|"):
            if SEP_ROW_RE.match(line):
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            continue
        else:
            flush_table()

        if line.strip() == "":
            doc.add_paragraph()
            continue

        if RULE_RE.match(line):
            para = doc.add_paragraph()
            add_border(para, 'bottom')
            continue

        m = HEADING_RE.match(line)
        if m:
            lvl = len(m.group(1))
            heading = doc.add_heading(m.group(2), level=lvl)
            if lvl == 1:
                add_border(heading, 'bottom', color='2E5496', sz='12', space='4')
            if toc and not toc_done:
                add_table_of_contents(doc)
                toc_done = True
            continue

        m = BULLET_RE.match(line)
        if m:
            para = doc.add_paragraph(style='List Bullet')
            add_inline(para, m.group(1))
            if quoted:
                para.paragraph_format.left_indent = Pt(36)
                shade_paragraph(para, Q_BG)
                add_border(para, 'left', color=Q_BAR, sz="24", space="10")
            continue

        para = doc.add_paragraph()
        add_inline(para, line)
        if quoted:
            para.paragraph_format.left_indent = Pt(14)
            para.paragraph_format.space_after = Pt(2)
            shade_paragraph(para, Q_BG)
            add_border(para, 'left', color=Q_BAR, sz="24", space="10")

    # offene Bloecke schliessen
    if in_code and code_buf:
        add_code_block(doc, code_buf, code_lang)
    flush_table()

    if toc:
        enable_update_fields(doc)
    doc.save(str(docx_path))
    print("OK\t%s" % docx_path)


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    flags = [a for a in sys.argv[1:] if a.startswith("--")]
    if len(args) < 2:
        sys.stderr.write("usage: build_docx.py <input.md> <output.docx> [--light] [--no-toc]\n")
        sys.exit(1)
    md_path = pathlib.Path(args[0])
    docx_path = pathlib.Path(args[1])
    theme = "light" if "--light" in flags else "dark"
    toc = "--no-toc" not in flags
    if not md_path.is_file():
        sys.stderr.write("Eingabedatei nicht gefunden: %s\n" % md_path)
        sys.exit(1)
    build(md_path, docx_path, theme=theme, toc=toc)


if __name__ == "__main__":
    main()
